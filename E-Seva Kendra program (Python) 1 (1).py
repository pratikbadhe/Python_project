from easygui import *
import sys
import Tkinter
import tkMessageBox
msgbox("Bonjour Monsiuer/Mademosielle")
while 1:
 a = choicebox("Apply for","E-Seva Kendra",choices=["Passport","Aadhar Card"])
 if a == "Passport":
  a0 = ccbox("You have opted for Passport Services of India")
  a1 = enterbox("Enter applicant's full name(in caps)")
  a2 = enterbox("Birth Date")
  a3 = enterbox("City")
  a4 = enterbox("District")
  a5 = enterbox("State")
  #a22= fileopenbox("upload your photo")
  a6 = choicebox("Gender",choices= ["Male","Female","Other"])
  a7 = choicebox("Are you married?","Marital status",choices = ["Yes","No"])
  a8 = enterbox("Enter Pan card number 'if available' or enter '-' if not available")
  a9 = enterbox("Enter Aadhar card number")
  a10 = enterbox("Enter Voter id 'if available' or enter '-' if not available")
  a11 = choicebox("Employment type",choices = ["Student","Business","Government sector","Employee","Housewife","none"])
  a12 = choicebox("Is either of your parent (In case of minor i.e age below 18) a government servant",choices = ["Yes","No"])
  a13 = choicebox("Education qualification",choices =["<10","Matric","10+2","10+2[+]Graduation","10+2[+]Graduation[+]Post Graduation","Diploma"])
  #a14 = choicebox("Provide atleast one's details",choices = ["Father","Mother","Legal guardian"])
  a14 = choicebox("Is any of your family member serving in Indian armed forces",choices = ["Yes","No"])
  a15 = enterbox("Enter your relation with the person who was a freedom fighter if any")
  c1 = msgbox("I solemly swear that the above mentioned information is true. \n")
  c2 = msgbox("\t\t Yours faithfully, \n ______")
  from docx import Document
  document=Document()
  #document.add_picture("a22")
  #paragraph = document.add_paragraph('Lorem ipsum dolor sit amet.')
  document.save('Passport.docx')
  f = open('Passport.docx', 'rb')
  document = Document(f)
  paragraph=document.add_paragraph("Name = "+ a1 +"\nBirth Date = "+a2 +"\nCity = "+ a3 +"\nDistrict = "+ a4 +"\nState = "+ a5+ "\nGender = "+ a6+ "\nMartial status = "+a7 +"\nPan card number = "+ a8 + "\nAadhar number = "+ a9 + "\nVoter ID = "+ a10 +"\nEmployment Type = "+ a11 + "\nIs either parent a government servant = "+ a12 + "\nEducational Qualification = " + a13 +"\nIs anyone of your family serve in Indian Armed Force = "+ a14 +"\nRelation With Soldier if any = "+ "\nI solemly swear that the above mentioned information is true." +"\n\t\t Yours faithfully, \n ______" )
  document.save('Passport.docx')
  sys.exit(0)
 elif a == "Aadhar Card":
  b = ccbox("You have opted for UIDAI (Unique Identification Authority of India) service")
  b1 = enterbox('''Enter applicant's full name(in caps)''')
  b2 = enterbox("Birth Date")
  b3 = enterbox("City")
  b4 = enterbox("District")
  b5 = enterbox("State")
  b6 = choicebox("Gender",choices= ["Male","Female","Other"])
  b7 = choicebox("Are you married?","Marital status",choices = ["Yes","No"])
  b8 = enterbox('''Enter Pan card number 'if available' or enter '-' if not available''')
  b9 = enterbox('''Please enter any of your father/mother/guardian’s Aadhar card number''')
  c = msgbox("I confirm that I have been residing in India for at least 182 days in the preceding 12 months & information (including biometrics) provided by me to the UIDAI is my own and is true, correct and accurate. I am aware that my information including biometrics  will be used for generation of Aadhaar and authentication. I understand that my identity  information (except core biometric) may be provided to an agency only with my consent  during authentication or as per the provisions of the Aadhaar Act. I have a right to access my  identity information (except core biometrics) following the procedure laid down by UIDAI.  Verifier’s Stamp and Signature: (Verifier must put his/her Name, if stamp is not available)  Applicant’s signature/Thumbprint  ---------------------------------------------------------------------------------------------------------------------------  ---------------------------------------------- To be filled by the Enrolment Agency only : Date & time  of Enrolment: ---------------------------------------------------- “(Note: Incase of minor, the signature  will be done by parent/guardian. Incase of incapacitated person, the signature will be done by Legal Guardian of Incapacitated Person)")
  from docx import Document
  document=Document()
  #document.add_picture("a22")
  #paragraph = document.add_paragraph('Lorem ipsum dolor sit amet.')
  document.save('Aadhar.docx')
  f = open('Aadhar.docx', 'rb')
  document = Document(f)
  paragrah=document.add_paragraph("Name = "+ b1 +"\nBirth Date = "+ b2 +"\nCity = "+ b3 +"\nDistrict = "+ b4 +"\nState = "+ b5 + "\nGender = "+ b6 + "\nMartial status = "+ b7 +"\nPan card number = "+ b8 + "\nAadhar number of parent/Guardian= "+ b9 + "\n)
  document.save('Aadhar.docx')
  sys.exit(0)
 else:
  sys.exit(0)
     
